"""
Generates Report_1 (with embedded images), Report_2 (expanded),
Report_3 (code documentation), and INFO_REPORT.txt.
Run: ./venv/bin/python generate_reports.py  (from final/ directory)
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ─────────────────────────────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────────────────────────────

def new_doc():
    doc = Document()
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = Pt(12)
    for sec in doc.sections:
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(3.0)
        sec.right_margin  = Cm(2.5)
    return doc

def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.runs[0]
    run.font.color.rgb = RGBColor(0x1a, 0x4f, 0x8a)
    run.font.name = 'Times New Roman'
    return p

def body(doc, text, bold=False, italic=False, sa=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(sa)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold   = bold
    run.italic = italic
    return p

def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(10)
    run = p.add_run(text)
    run.font.name  = 'Times New Roman'
    run.font.size  = Pt(10)
    run.italic     = True

def insert_image(doc, path, width_cm=14):
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        run = p.add_run()
        run.add_picture(path, width=Cm(width_cm))
    else:
        body(doc, f"[IMAGE NOT FOUND: {path}]")

def toc_line(doc, text, page, level=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.left_indent  = Cm(level * 0.7)
    pPr = p._p.get_or_add_pPr()
    tabs_el = OxmlElement('w:tabs')
    tab_el  = OxmlElement('w:tab')
    tab_el.set(qn('w:val'),    'right')
    tab_el.set(qn('w:leader'), 'dot')
    tab_el.set(qn('w:pos'),    '7920')
    tabs_el.append(tab_el)
    pPr.append(tabs_el)
    run = p.add_run(text + '\t' + str(page))
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12 if level == 0 else 11)
    run.bold = (level == 0)

def add_toc(doc):
    heading(doc, "Table of Contents")
    toc_line(doc, "Abstract",                               3,  0)
    toc_line(doc, "1.  Introduction",                       3,  0)
    toc_line(doc, "1.1  Dataset",                           4,  1)
    toc_line(doc, "1.2  Formal Problem Statement",          5,  1)
    toc_line(doc, "1.3  Key Challenges",                    5,  1)
    toc_line(doc, "2.  Related Work",                       6,  0)
    toc_line(doc, "3.  Proposed Approach",                  7,  0)
    toc_line(doc, "3.1  Naïve Bayes Collaborative Filtering", 7, 1)
    toc_line(doc, "3.2  Log-Space Computation",             7,  1)
    toc_line(doc, "3.3  Temporal Decay Weighting",          8,  1)
    toc_line(doc, "3.4  Top-K Neighborhood Selection",      8,  1)
    toc_line(doc, "3.5  Hybrid Combination",                9,  1)
    toc_line(doc, "3.6  Confidence Scoring",               10,  1)
    toc_line(doc, "4.  Experimental Setup",                10,  0)
    toc_line(doc, "4.1  Train / Test Split",               10,  1)
    toc_line(doc, "4.2  Evaluation Subsets",               10,  1)
    toc_line(doc, "4.3  Hyperparameters",                  11,  1)
    toc_line(doc, "4.4  Evaluation Metrics",               11,  1)
    toc_line(doc, "5.  Performance Analysis",              12,  0)
    toc_line(doc, "5.1  Overall Results",                  12,  1)
    toc_line(doc, "5.2  User Activity Group Analysis",     14,  1)
    toc_line(doc, "5.3  Confidence Analysis",              16,  1)
    toc_line(doc, "5.4  Alpha Sensitivity",                17,  1)
    toc_line(doc, "5.5  Latency and Efficiency",           18,  1)
    toc_line(doc, "6.  Conclusion",                        20,  0)
    toc_line(doc, "References",                            21,  0)
    doc.add_page_break()

def cover(doc, title, report_no, supervisor, students):
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("RECOMMENDER SYSTEMS")
    r.bold = True; r.font.size = Pt(18); r.font.name = 'Times New Roman'

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(title)
    r2.bold = True; r2.font.size = Pt(14); r2.font.name = 'Times New Roman'

    doc.add_paragraph()
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(report_no)
    r3.font.size = Pt(13); r3.font.name = 'Times New Roman'

    for _ in range(3):
        doc.add_paragraph()
    for line in [f"Supervisor: {supervisor}"] + students:
        px = doc.add_paragraph()
        px.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rx = px.add_run(line)
        rx.font.size = Pt(12); rx.font.name = 'Times New Roman'
    doc.add_page_break()

def simple_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        c = table.rows[0].cells[j]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.name = 'Times New Roman'
        c.paragraphs[0].runs[0].font.size = Pt(10)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = table.rows[i + 1].cells[j]
            c.text = str(val)
            c.paragraphs[0].runs[0].font.name = 'Times New Roman'
            c.paragraphs[0].runs[0].font.size = Pt(10)
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if col_widths:
        for row in table.rows:
            for j, cell in enumerate(row.cells):
                cell.width = Cm(col_widths[j])
    doc.add_paragraph()

SUPERVISOR = "Gamze USLU, Ph.D."
STUDENTS   = ["Ömer MANAV – 210209003", "Kaan Bahadır SAĞRA – 210201011"]
OUT        = "outputs/"

# ══════════════════════════════════════════════════════════════════════
# REPORT 1  —  Full Academic Paper with Embedded Figures
# ══════════════════════════════════════════════════════════════════════

doc1 = new_doc()
cover(doc1,
      "A Collaborative Filtering Approach Based on Naïve Bayes Classifier",
      "Report 1 — Final Project",
      SUPERVISOR, STUDENTS)

add_toc(doc1)

# ── Abstract ──────────────────────────────────────────────────────────
heading(doc1, "Abstract")
body(doc1,
    "Rating prediction in collaborative filtering (CF) is challenging due to "
    "high matrix sparsity and temporal preference drift. This project implements "
    "and extends a Naïve Bayes Collaborative Filtering (NBCF) system evaluated "
    "on the MovieLens-1M benchmark. The midterm version introduced a hybrid "
    "user-based / item-based NBCF with exponential temporal decay weighting "
    "and entropy-based confidence scoring. The final version adds three "
    "substantial improvements: (1) Top-K neighborhood selection that restricts "
    "evidence to the K=30 most similar neighbors ranked by co-rater overlap, "
    "(2) RMSE as a complementary evaluation metric alongside MAE, and "
    "(3) stratified evaluation across three user activity groups to expose "
    "cold-start behavior. The final model (TopK+Temporal) achieves MAE = 0.798 "
    "and RMSE = 1.126 on the held-out test set, with a 59% latency reduction "
    "compared to the full-neighbourhood temporal baseline (34.8 ms → 14.3 ms "
    "per prediction). Confidence-filtered predictions covering 29.3% of the "
    "test set achieve a filtered MAE of 0.706, confirming that the model's "
    "uncertainty estimates are informative. An alpha sensitivity analysis shows "
    "that the temporal variant is robust to the Laplace smoothing parameter.")

# ── 1. Introduction ───────────────────────────────────────────────────
heading(doc1, "1. Introduction")
body(doc1,
    "Online platforms such as streaming services and e-commerce sites offer "
    "millions of items, making manual discovery infeasible. Recommender systems "
    "address this by predicting user preferences from historical behavior. "
    "Collaborative Filtering (CF), the dominant paradigm, assumes that users "
    "who agreed in the past will agree again in the future, requiring no item "
    "metadata and scaling naturally with usage data.")
body(doc1,
    "This project frames rating prediction as a five-class classification task "
    "and applies the Naïve Bayes Collaborative Filtering algorithm. NBCF "
    "computes the posterior probability of each rating class using Bayes' "
    "theorem, treating past ratings as conditionally independent features. "
    "We evaluate on MovieLens-1M — 1,000,209 ratings by 6,040 users on "
    "3,701 movies.")
body(doc1,
    "Core contributions: (i) hybrid UB/IB NBCF with log-space computation and "
    "Laplace smoothing; (ii) exponential temporal decay; (iii) Top-K "
    "neighborhood selection; (iv) max-probability confidence measure with "
    "coverage–MAE trade-off analysis; (v) stratified evaluation by user "
    "activity group.")

heading(doc1, "1.1 Dataset", level=2)
body(doc1, "The MovieLens-1M dataset properties are summarised below.")
simple_table(doc1,
    ["Property", "Value"],
    [["Number of users", "6,040"], ["Number of movies", "3,701"],
     ["Total ratings",   "1,000,209"], ["Rating scale", "1–5 (integer)"],
     ["Matrix sparsity", "95.53 %"], ["Timestamp range", "2000–2003"]],
    col_widths=[6, 5])

body(doc1,
    "Table 1 below shows a small extract of the actual rating matrix from the "
    "dataset, illustrating the typical sparsity pattern.")
insert_image(doc1, OUT + "table1_rating_matrix.png", width_cm=14)
caption(doc1, "Table 1. Running example of the rating matrix (ML-1M). '•' denotes a missing entry.")

heading(doc1, "1.2 Formal Problem Statement", level=2)
body(doc1,
    "Let U and I be the sets of users and items. The rating matrix R is "
    "partially observed. The goal is to learn f̂(u, i) → ŷ ∈ {1,2,3,4,5}. "
    "Quality is measured with MAE = (1/|T|) Σ|r−ŷ| and "
    "RMSE = √((1/|T|) Σ(r−ŷ)²), where T is the test set.")

heading(doc1, "1.3 Key Challenges", level=2)
bullet(doc1, "Data sparsity — 95.53% of entries are missing.")
bullet(doc1, "Preference drift — user tastes change over time.")
bullet(doc1, "Scalability — popular items can have thousands of co-raters.")
bullet(doc1, "Uncertainty — the model should indicate when it is unsure.")
doc1.add_paragraph()

# ── 2. Related Work ───────────────────────────────────────────────────
heading(doc1, "2. Related Work")
body(doc1,
    "Breese et al. (1998) systematically compared memory-based CF variants and "
    "established MAE as a standard metric. Herlocker et al. (2004) provided a "
    "comprehensive evaluation framework still in common use today.")
body(doc1,
    "The Naïve Bayes formulation of CF was introduced by Si and Jin (2003), "
    "who showed that treating rating prediction as classification under the "
    "conditional independence assumption yields competitive accuracy while "
    "remaining interpretable. Laplace smoothing and log-space computation "
    "address zero-probability and numerical underflow in sparse settings.")
body(doc1,
    "Temporal dynamics were incorporated into CF by Koren (2009) through "
    "TimeSVD++, which models item popularity drift and user taste evolution. "
    "Our exponential decay weighting is a lighter, training-free alternative.")
body(doc1,
    "Neighborhood selection in CF was studied by Sarwar et al. (2001), who "
    "showed that limiting item-based CF to the K most similar neighbors improves "
    "both accuracy and efficiency. We adapt this to NBCF using co-rater overlap "
    "as the similarity measure.")
body(doc1,
    "Confidence and uncertainty in recommender systems were studied by "
    "Hernandez-del-Olmo and Gaudioso (2008) through coverage–precision curves, "
    "analogous to our Figure 8 confidence threshold sweep.")

# ── 3. Proposed Approach ──────────────────────────────────────────────
heading(doc1, "3. Proposed Approach")

heading(doc1, "3.1 Naïve Bayes Collaborative Filtering", level=2)
body(doc1,
    "We find y ∈ {1,2,3,4,5} that maximizes the posterior. Under the Naïve "
    "Bayes assumption, past ratings are conditionally independent given y. "
    "The user-based (UB) variant gives:")
body(doc1,
    "    P(y | u, i) ∝ P(y | i) × Π_{j ∈ I_u} P(r_{u,j}=k_j | y, i)", italic=True)
body(doc1,
    "where P(y | i) is the Laplace-smoothed item prior. The item-based (IB) "
    "variant uses user priors and co-raters of item i instead.")

body(doc1, "Tables 2 and 3 show the prior probabilities and user-based likelihoods "
    "computed for a concrete example from the dataset.")
insert_image(doc1, OUT + "table2_3_priors_likelihood.png", width_cm=14)
caption(doc1, "Tables 2–3. Prior probabilities and user-based likelihoods for a running example.")

heading(doc1, "3.2 Log-Space Computation", level=2)
body(doc1,
    "Multiplying many small probabilities causes underflow. We use log-space "
    "accumulation and the log-sum-exp trick for numerically stable normalisation:")
body(doc1,
    "    log_s(y) = log P(y|i) + Σ_j log P(r_{u,j}=k | y)", italic=True)
body(doc1,
    "    P(y) = exp(log_s(y) − max_y log_s) / Σ_y exp(log_s(y) − max_y log_s)", italic=True)

heading(doc1, "3.3 Temporal Decay Weighting (Midterm Contribution)", level=2)
body(doc1,
    "We weight each past rating by its recency: "
    "w(u, j) = exp(−λ × (t_max − t(u,j))), λ = 1×10⁻⁷. "
    "A rating from three years ago receives weight ≈ 0.09. The weight "
    "multiplies each log-likelihood term inside the score accumulation loop.")

heading(doc1, "3.4 Top-K Neighborhood Selection (Final Contribution)", level=2)
body(doc1,
    "Original NBCF uses all past items (UB) or all co-raters (IB). For popular "
    "items, the IB side accumulates up to 2,720 co-raters, most sharing few "
    "items with the target user and contributing near-uniform likelihoods — "
    "noise rather than signal.")
body(doc1,
    "Top-K selection ranks candidate neighbors by co-rater overlap and retains "
    "the top K = 30. For UB: score(j, i) = |users(i) ∩ users(j)|. "
    "This focuses evidence on the most relevant neighbors and caps the inner "
    "loop at O(K), yielding a 59% latency improvement.")

body(doc1, "Table 4 shows item-based likelihoods, and Table 5 compares the full "
    "score distributions for all nine model variants on a concrete example.")
insert_image(doc1, OUT + "table4_5_ib_scores.png", width_cm=14)
caption(doc1, "Tables 4–5. Item-based likelihoods and classification scores for all variants "
              "(* marks the predicted class).")

heading(doc1, "3.5 Hybrid Combination", level=2)
body(doc1,
    "UB and IB scores are combined geometrically with adaptive weights: "
    "w_UB = 1/(1+|I_u|), w_IB = 1/(1+|U_i|), "
    "P_hybrid(y) ∝ P_UB(y)^w_UB × P_IB(y)^w_IB.")

heading(doc1, "3.6 Confidence Scoring", level=2)
body(doc1,
    "We use the maximum predicted probability as the confidence score: "
    "confidence = max_y P(y|u,i). Range: [0.2, 1.0] for five classes. "
    "The midterm entropy-based confidence degenerated to near-zero under the "
    "geometric hybrid (exponents w ≈ 0.04 push all probabilities toward 1). "
    "Max-probability avoids this by operating on the UB-Temporal component "
    "scores before geometric mixing.")

# ── 4. Experimental Setup ─────────────────────────────────────────────
heading(doc1, "4. Experimental Setup")

heading(doc1, "4.1 Train / Test Split", level=2)
body(doc1,
    "The 1,000,209 ratings are split 80/20 (random_state=42): "
    "800,167 training, 200,042 test ratings.")

heading(doc1, "4.2 Evaluation Subsets", level=2)
bullet(doc1,
    "Main (n=1,000): test samples from users with 20–30 training ratings. "
    "Identical to midterm for direct comparison.")
bullet(doc1,
    "User-group (300 each): Light (10–30), Medium (31–100), Heavy (>100 ratings).")
bullet(doc1,
    "Alpha sensitivity (n=200): 5 values of α on the main subset.")

heading(doc1, "4.3 Hyperparameters", level=2)
simple_table(doc1,
    ["Parameter", "Value", "Description"],
    [["α (Laplace smoothing)", "0.01",   "Prevents zero probabilities"],
     ["λ (decay rate)",        "1×10⁻⁷", "Temporal exponential decay"],
     ["K (neighborhood)",      "30",     "Max neighbors in Top-K"],
     ["τ (confidence)",        "0.45",   "Min confidence for filtering"]],
    col_widths=[5, 3, 7])

heading(doc1, "4.4 Evaluation Metrics", level=2)
bullet(doc1, "MAE — Mean Absolute Error.")
bullet(doc1, "RMSE — Root Mean Squared Error; penalizes large errors more.")
bullet(doc1, "Macro Precision / Recall / F1 — per-class then averaged equally.")
bullet(doc1, "Binary Confusion Matrix — Like/Dislike (rating ≥ 4 = Like).")
bullet(doc1, "Coverage — fraction of predictions above confidence threshold τ.")
bullet(doc1, "Filtered MAE — MAE on high-confidence predictions only.")
doc1.add_paragraph()

# ── 5. Performance Analysis ───────────────────────────────────────────
heading(doc1, "5. Performance Analysis")

heading(doc1, "5.1 Overall Results", level=2)
body(doc1,
    "The table below summarises all five NBCF variants. Key observations: "
    "UB (MAE=0.861) strongly outperforms IB (MAE=2.335) due to noise from "
    "uninformative co-raters on popular items. Temporal decay reduces "
    "Hybrid MAE by 28% (1.102 → 0.792). TopK+Temporal matches Temporal "
    "accuracy at 59% lower latency.")
simple_table(doc1,
    ["Method", "MAE", "RMSE", "Precision", "Recall", "F1", "Coverage"],
    [["UB",                    "0.8610","1.2518","—","—","—","—"],
     ["IB",                    "2.3350","2.6384","—","—","—","—"],
     ["Hybrid",                "1.1020","1.4953","0.2752","0.2974","0.2706","—"],
     ["Hybrid + Temporal",     "0.7920","1.1234","0.3720","0.2761","0.2792","—"],
     ["TopK + Temporal [NEW]", "0.7980","1.1261","0.3638","0.2658","0.2656","29.3%"]],
    col_widths=[5, 2.2, 2.2, 2.3, 2.2, 2.2, 2.2])
caption(doc1, "Table 7. Comprehensive evaluation metrics for all NBCF variants (n = 1,000).")

insert_image(doc1, OUT + "table7_full_metrics.png", width_cm=14)
caption(doc1, "Table 7 (visual). Full metrics summary generated by the evaluation script.")

body(doc1, "Figure 2 visualises the MAE and RMSE for all five methods side by side.")
insert_image(doc1, OUT + "figure2_mae_rmse_comparison.png", width_cm=14)
caption(doc1, "Figure 2. MAE (left) and RMSE (right) comparison across all NBCF variants.")

body(doc1, "Figure 3 shows per-class F1 scores for the three main variants.")
insert_image(doc1, OUT + "figure3_per_class_f1.png", width_cm=13)
caption(doc1, "Figure 3. Per-class F1 score by rating value.")

body(doc1, "Figure 5 compares the predicted rating distributions and per-class MAE.")
insert_image(doc1, OUT + "figure5_rating_distribution.png", width_cm=14)
caption(doc1, "Figure 5. Predicted vs actual rating distribution (left) and MAE per rating class (right).")

body(doc1, "Figure 1 shows the confusion matrices for the best model (TopK+Temporal).")
insert_image(doc1, OUT + "figure1_confusion_matrices.png", width_cm=14)
caption(doc1, "Figure 1. Confusion matrices for TopK+Temporal NBCF: 5-class and binary (Like/Dislike).")

heading(doc1, "5.2 User Activity Group Analysis", level=2)
body(doc1,
    "Figure 7 and Table 6 break down performance by user activity. "
    "Light users show no difference between models (K=30 rarely filters "
    "their ≤30-item history). Medium users benefit most: TopK reduces MAE "
    "by 6.2% (0.803 → 0.753). Heavy users (>100 ratings) lose slightly "
    "(0.740 → 0.777) because K=30 truncates genuinely useful history. "
    "The overall cold-start trend — heavier users have lower error — "
    "is clearly visible in both metrics.")
insert_image(doc1, OUT + "figure7_user_groups.png", width_cm=14)
caption(doc1, "Figure 7. MAE (left) and RMSE (right) by user activity group.")

insert_image(doc1, OUT + "table6_user_groups.png", width_cm=14)
caption(doc1, "Table 6. Numeric breakdown of performance by user activity group.")

heading(doc1, "5.3 Confidence Analysis", level=2)
body(doc1,
    "Figure 8 sweeps the confidence threshold τ from 0.20 to 0.95. "
    "At τ=0.45, coverage is 29.3% with filtered MAE = 0.706, an 11.5% "
    "improvement over the unfiltered MAE of 0.798. As τ rises, coverage "
    "falls and filtered MAE continues to fall, confirming that the "
    "confidence score is genuinely informative. Figure 4 shows the "
    "confidence distribution and its correlation with MAE per bucket.")
insert_image(doc1, OUT + "figure4_confidence.png", width_cm=14)
caption(doc1, "Figure 4. Confidence distribution (left) and MAE vs confidence bucket (right).")

insert_image(doc1, OUT + "figure8_confidence_sweep.png", width_cm=13)
caption(doc1, "Figure 8. Coverage–MAE trade-off as confidence threshold is varied.")

heading(doc1, "5.4 Alpha Sensitivity", level=2)
body(doc1,
    "Figure 9 plots MAE against five values of the Laplace smoothing "
    "parameter α. UB (no temporal) has a sweet spot at α ≈ 0.1 "
    "(MAE=0.845). UB+Temporal is nearly flat across all α ≥ 0.01 "
    "(MAE=0.855), showing that temporal weighting absorbs the impact of "
    "the smoothing choice and makes the model robust to this hyperparameter.")
insert_image(doc1, OUT + "figure9_alpha_sensitivity.png", width_cm=12)
caption(doc1, "Figure 9. MAE vs Laplace smoothing parameter α (log scale).")

heading(doc1, "5.5 Latency and Efficiency", level=2)
body(doc1,
    "Figure 6 compares the latency distributions and throughput scaling. "
    "TopK+Temporal is 2.43× faster than full Temporal because the inner "
    "scoring loop is capped at K=30 instead of the full co-rater count "
    "(median 100, max 2,720 per item).")
simple_table(doc1,
    ["Model", "Avg Latency", "p95 Latency", "Speedup vs Temporal"],
    [["Hybrid",            "28.3 ms","70.0 ms","—"],
     ["Hybrid+Temporal",   "34.8 ms","87.9 ms","baseline"],
     ["TopK+Temporal [NEW]","14.3 ms","35.1 ms","× 2.43"]],
    col_widths=[5.5, 3, 3, 4.5])
insert_image(doc1, OUT + "figure6_sustainability.png", width_cm=14)
caption(doc1, "Figure 6. Latency distribution (left) and throughput vs user history depth (right).")

# ── 6. Conclusion ─────────────────────────────────────────────────────
heading(doc1, "6. Conclusion")
body(doc1,
    "This project demonstrates that NBCF can be extended incrementally while "
    "remaining interpretable and training-free. The midterm version established "
    "a working hybrid NBCF with temporal decay and confidence scoring. The "
    "final version corrects the confidence degeneracy, introduces Top-K "
    "neighborhood selection, adds RMSE, and reveals cold-start behavior "
    "through stratified evaluation.")
body(doc1,
    "The most practically significant result is that Top-K selection reduces "
    "prediction latency by 59% while maintaining near-identical overall "
    "accuracy and improving accuracy by 6.2% for medium-activity users. "
    "Confidence filtering shows that high-confidence predictions have 11.5% "
    "lower MAE, confirming that uncertainty estimates are informative.")
body(doc1,
    "Future directions include adaptive K per user-activity level, "
    "mean-centering to correct for user rating bias, and cross-validation "
    "for more stable metric estimates.")

# ── 7. References ─────────────────────────────────────────────────────
heading(doc1, "References")
for ref in [
    "Breese, J. S., Heckerman, D., & Kadie, C. (1998). Empirical analysis of predictive "
    "algorithms for collaborative filtering. Proc. 14th UAI (pp. 43–52). AUAI.",
    "Herlocker, J. L., Konstan, J. A., Terveen, L. G., & Riedl, J. T. (2004). Evaluating "
    "collaborative filtering recommender systems. ACM TOIS, 22(1), 5–53.",
    "Hernandez-del-Olmo, F., & Gaudioso, E. (2008). Evaluation of recommender systems: "
    "A new approach. Expert Systems with Applications, 35(3), 790–804.",
    "Koren, Y. (2009). Collaborative filtering with temporal dynamics. Proc. 15th KDD "
    "(pp. 447–456). ACM.",
    "Harper, F. M., & Konstan, J. A. (2015). The MovieLens datasets: History and context. "
    "ACM TIIS, 5(4), 19.",
    "Sarwar, B., Karypis, G., Konstan, J., & Riedl, J. (2001). Item-based collaborative "
    "filtering recommendation algorithms. Proc. 10th WWW (pp. 285–295). ACM.",
    "Si, L., & Jin, R. (2003). Flexible mixture model for collaborative filtering. "
    "Proc. 20th ICML (pp. 704–711). ICML.",
]:
    p = doc1.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(ref)
    run.font.name = 'Times New Roman'; run.font.size = Pt(11)

doc1.save('Report_1_Final.docx')
print("✓ Report_1_Final.docx")


# ══════════════════════════════════════════════════════════════════════
# REPORT 2  —  Expanded Midterm vs Final Differences
# ══════════════════════════════════════════════════════════════════════

doc2 = new_doc()
cover(doc2,
      "A Collaborative Filtering Approach Based on Naïve Bayes Classifier",
      "Report 2 — Midterm vs. Final: Summary of Differences",
      SUPERVISOR, STUDENTS)

heading(doc2, "Overview")
body(doc2,
    "The final submission extends the midterm NBCF implementation with two algorithmic "
    "improvements and three evaluation additions. The midterm's temporal hybrid achieved "
    "MAE = 0.792 but had a non-functional confidence system (0% coverage) and evaluated "
    "only a single narrow user group. The final version addresses both issues and adds "
    "RMSE, stratified user-group analysis, and an alpha sensitivity study.", sa=6)

heading(doc2, "Algorithm Changes")
simple_table(doc2,
    ["Change", "Description", "Impact"],
    [
        ["Top-K Neighbor\nSelection (NEW)",
         "Candidates ranked by co-rater overlap\n(|users(i)∩users(j)| for UB;\n"
         "|items(u)∩items(v)| for IB).\nTop K=30 retained; rest discarded.",
         "Latency: 34.8ms → 14.3ms (−59%)\n"
         "MAE medium users: 0.803 → 0.753 (−6.2%)\n"
         "MAE heavy users: 0.740 → 0.777 (+slight)"],
        ["Confidence Score\nFixed (UPDATED)",
         "Entropy-based measure replaced with\n"
         "max-probability: conf = max P(y|u,i)\n"
         "applied to UB-Temporal scores before\ngeometric hybrid mixing.",
         "Coverage: 0% → 29.3% (at τ=0.45)\n"
         "Filtered MAE: 0.706 (−11.5% vs overall)"],
    ],
    col_widths=[3.8, 6.2, 5.5])

heading(doc2, "Evaluation Changes")
simple_table(doc2,
    ["Addition", "Midterm", "Final"],
    [["RMSE metric",         "Not reported",              "Added alongside MAE"],
     ["User group eval",     "20–30 rating users only",   "Light / Medium / Heavy (300 each)"],
     ["Alpha sensitivity",   "α=0.01 fixed",              "5 values: 0.001 – 1.0"],
     ["Confidence sweep",    "Single threshold τ=0.30",   "Sweep τ=0.20 to 0.95 (Figure 8)"],
     ["New figures/tables",  "6 figures, 5 tables",       "9 figures, 7 tables"]],
    col_widths=[4.5, 5, 6])

heading(doc2, "Metric Comparison")
simple_table(doc2,
    ["Model", "MAE", "RMSE", "Coverage", "Latency"],
    [["UB",                   "0.861","1.252","—",    "—"],
     ["IB",                   "2.335","2.638","—",    "—"],
     ["Hybrid",               "1.102","1.495","—",    "28.3 ms"],
     ["Hybrid + Temporal",    "0.792","1.123","0%",   "34.8 ms"],
     ["TopK + Temporal [NEW]","0.798","1.126","29.3%","14.3 ms"]],
    col_widths=[4.5, 2.2, 2.2, 2.5, 2.8])

doc2.save('Report_2_Final.docx')
print("✓ Report_2_Final.docx")


# ══════════════════════════════════════════════════════════════════════
# REPORT 3  —  Line-by-Line Code and Output Documentation
# ══════════════════════════════════════════════════════════════════════

doc3 = new_doc()
cover(doc3,
      "A Collaborative Filtering Approach Based on Naïve Bayes Classifier",
      "Report 3 — Code Documentation",
      SUPERVISOR, STUDENTS)

def code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.left_indent  = Cm(1.0)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1a, 0x4f, 0x8a)

heading(doc3, "Overview")
body(doc3,
    "This document explains every section of nbcf_final.py line by line and "
    "describes the meaning of each section of the output. The script has nine "
    "numbered sections: Sections 1–4 define data and algorithms; "
    "Sections 5–7 run evaluations; Sections 8–9 produce all figures and tables.")
body(doc3,
    "New additions in the final version (not present in the midterm) are marked "
    "with [NEW] in the section headings.")

heading(doc3, "Section 1 — Imports and Global Constants")
heading(doc3, "1.1 Library Imports", level=2)
for line, expl in [
    ("import pandas as pd",
     "Reads the .dat dataset file and provides DataFrame operations."),
    ("import numpy as np",
     "Numerical operations: mean, sqrt, percentile, array arithmetic."),
    ("import matplotlib.pyplot as plt",
     "Creates all figures and saves them as PNG files."),
    ("import seaborn as sns",
     "Draws the confusion-matrix heatmaps with color gradients."),
    ("import math",
     "math.log() and math.exp() for the log-space probability calculations."),
    ("import time",
     "time.perf_counter() times each prediction to nanosecond precision."),
    ("from collections import defaultdict",
     "Nested dictionaries that return empty dicts for missing keys — no KeyError."),
    ("from sklearn.model_selection import train_test_split",
     "Splits all ratings 80/20 with a fixed seed for reproducibility."),
    ("from sklearn.metrics import confusion_matrix, precision_score, ...",
     "Standard classification metrics: confusion matrix, precision, recall, F1."),
]:
    code(doc3, line)
    body(doc3, expl, sa=4)

heading(doc3, "1.2 Global Constants", level=2)
for line, expl in [
    ("R_values = [1, 2, 3, 4, 5]",
     "The five rating classes. Every probability distribution is computed over these five values."),
    ("ALPHA = 0.01",
     "Laplace smoothing parameter. Prevents any probability from being exactly zero when a "
     "rating class has never appeared for a given item or user."),
    ("DECAY_LAMBDA = 1e-7",
     "Exponential decay rate. At this value, a rating made 3 years ago (≈ 9.5×10^7 seconds) "
     "receives weight exp(−1×10⁻⁷ × 9.5×10⁷) ≈ 0.09 — about 9% of a recent rating."),
    ("TOP_K = 30",
     "[NEW] Maximum number of neighbors used per prediction. Caps the inner scoring loop "
     "at 30 iterations regardless of how many past items or co-raters exist."),
]:
    code(doc3, line)
    body(doc3, expl, sa=4)

heading(doc3, "Section 2 — Dataset Loading and Preparation")
for line, expl in [
    ("ratings_df = pd.read_csv('ml-1m/ratings.dat', sep='::', ...)",
     "Reads the dataset. sep='::' is the MovieLens-1M delimiter. encoding='latin-1' "
     "handles extended characters in movie titles."),
    ("train_df, test_df = train_test_split(..., test_size=0.2, random_state=42)",
     "80% training, 20% test. random_state=42 gives an identical split on every run."),
    ("user_r[row['user_id']][row['item_id']] = row['rating']",
     "user_r[u][i] holds the rating of user u on item i. Enables O(1) lookup."),
    ("user_ts[row['user_id']][row['item_id']] = row['timestamp']",
     "Stores when each rating was made. Used by temporal_weight() to compute age."),
    ("item_users_set = {i: frozenset(...) for i, u_dict in item_r.items()}",
     "[NEW] Precomputes the set of users who rated each item. frozenset supports fast "
     "set-intersection (& operator), which is used to rank Top-K neighbors by overlap. "
     "Precomputing avoids rebuilding the set on every prediction call."),
]:
    code(doc3, line)
    body(doc3, expl, sa=4)

heading(doc3, "Section 3 — Core Algorithm Functions (Unchanged from Midterm)")
heading(doc3, "3.1 Prior Probabilities", level=2)
code(doc3, "def calc_item_prior(i, y, a=ALPHA):")
code(doc3, "    return (count_y + a) / (len(ratings) + NUM_R * a)")
body(doc3,
    "Estimates P(y | i): the probability that a random rating of item i equals y. "
    "Laplace smoothing adds a (numerator) and NUM_R × a (denominator) so the "
    "probability is never zero even for unseen rating classes.", sa=4)
code(doc3, "def calc_user_prior(u, y, a=ALPHA):")
body(doc3, "Same formula applied to user u instead of item i, estimating P(y | u).", sa=4)

heading(doc3, "3.2 Likelihood Functions", level=2)
code(doc3, "def calc_ub_likelihood(i, y, j, k, a=ALPHA):")
body(doc3,
    "Computes P(r_{u,j}=k | rating of item i = y). "
    "Step 1: find all users who gave item i rating y. "
    "Step 2: among those, count how many gave item j rating k (match) "
    "and how many rated item j at all (total). "
    "Step 3: return (match + a) / (total + NUM_R × a).", sa=4)
code(doc3, "def calc_ib_likelihood(u, y, v, k, a=ALPHA):")
body(doc3,
    "Computes P(co-rater v gives k | user u gives y to their items). "
    "Finds items user u rated as y, counts how many co-rater v also rated "
    "with rating k. Same Laplace formula applied.", sa=4)

heading(doc3, "3.3 Log-Space Scoring", level=2)
code(doc3, "def normalize_log_scores(log_scores_dict):")
body(doc3,
    "Log-sum-exp trick: subtract the maximum log-score before calling exp() "
    "to avoid overflow/underflow, then divide by the sum to normalize to a "
    "valid probability distribution that sums to 1.", sa=4)
code(doc3, "def ub_score(u, i, a=ALPHA):")
body(doc3,
    "For each class y, accumulates log P(y|i) + Σ_j log P(r_{u,j}=k|y) over "
    "ALL past items j of user u, then normalizes. The 'a' parameter allows "
    "alpha sensitivity analysis to override the default.", sa=4)
code(doc3, "def ib_score(u, i, a=ALPHA):")
body(doc3, "Same structure but iterates over ALL co-raters of item i instead.", sa=4)

heading(doc3, "3.4 Temporal Score Functions", level=2)
code(doc3, "def temporal_weight(user_id, item_id):")
code(doc3, "    return math.exp(-DECAY_LAMBDA * (global_max_ts - ts))")
body(doc3,
    "Computes the recency weight for a past rating. "
    "global_max_ts is the most recent timestamp in the training set. "
    "age = global_max_ts − t(u,j) is measured in seconds. "
    "If a rating's timestamp is not found (missing key), age defaults to 0 → weight 1.", sa=4)
code(doc3, "def ub_score_temporal(u, i, a=ALPHA):")
body(doc3,
    "Same as ub_score but multiplies each log-likelihood term by w(u,j) "
    "before adding it to log_s. Recent items contribute their full "
    "log-likelihood; older items contribute a fraction of it.", sa=4)

heading(doc3, "Section 4 — Top-K Neighborhood Selection [NEW]")
heading(doc3, "4.1 topk_items_for_ub", level=2)
code(doc3, "def topk_items_for_ub(u, i, K=TOP_K):")
code(doc3, "    i_set = item_users_set.get(i, frozenset())")
code(doc3, "    return sorted(past,")
code(doc3, "        key=lambda jk: len(i_set & item_users_set.get(jk[0], frozenset())),")
code(doc3, "        reverse=True)[:K]")
body(doc3,
    "If user u has ≤ K past items: returns all (no filtering). "
    "Otherwise: computes |users(i) ∩ users(j)| for each past item j. "
    "A high overlap means items i and j are rated by the same community, "
    "so j is a strong evidence source for predicting i. "
    "Sorting descending and slicing to K keeps only the most relevant.", sa=4)

heading(doc3, "4.2 topk_users_for_ib", level=2)
code(doc3, "def topk_users_for_ib(u, i, K=TOP_K):")
body(doc3,
    "Analogous to topk_items_for_ub. Ranks co-raters v of item i by "
    "|items(u) ∩ items(v)|. Co-raters who share more items with u "
    "give more informative IB likelihood estimates.", sa=4)

heading(doc3, "4.3 confidence_maxprob", level=2)
code(doc3, "def confidence_maxprob(scores_dict):")
code(doc3, "    return max(scores_dict.values())")
body(doc3,
    "Returns the probability of the most likely class. Range [0.2, 1.0] "
    "for five classes (0.2 = perfectly uniform, 1.0 = certain). "
    "Applied to UB-Temporal component scores before geometric mixing. "
    "The midterm entropy measure degenerated to ~0 for all predictions "
    "because the geometric hybrid exponents w_UB, w_IB ≈ 0.04 push all "
    "class probabilities toward 1, collapsing the distribution.", sa=4)

heading(doc3, "4.4 hybrid_topk_temporal — Final Model", level=2)
code(doc3, "def hybrid_topk_temporal(u, i, a=ALPHA):")
code(doc3, "    ub_s = ub_score_topk_temporal(u, i, a)")
code(doc3, "    ib_s = ib_score_topk_temporal(u, i, a)")
code(doc3, "    hybrid[y] = max(ub_s[y], 1e-300)**w_ub * max(ib_s[y], 1e-300)**w_ib")
body(doc3,
    "Calls the two Top-K temporal score functions then combines them with the "
    "same adaptive geometric weights as the original hybrid. "
    "1e-300 floor prevents log(0) errors.", sa=4)

heading(doc3, "Section 5 — Main Evaluation Loop")
for line, expl in [
    ("eval_df = test_df[test_df['user_id'].isin(small_users)].head(1000)",
     "1,000 test rows from users with 20–30 training ratings. Same subset as "
     "midterm to enable direct before/after comparison."),
    ("rmse_htk = np.sqrt(np.mean(np.array(errors_htk) ** 2))",
     "[NEW] RMSE: takes squared errors, averages them, takes the square root. "
     "Larger errors contribute disproportionately compared to MAE."),
    ("confidences_max.append(confidence_maxprob(ubtk_s))",
     "[NEW] Records max-probability confidence for each prediction using the "
     "UB Top-K Temporal scores."),
    ("cm_full = confusion_matrix(actuals, preds_htk, labels=R_values)",
     "Builds the 5×5 confusion matrix for the best model (TopK+Temporal). "
     "Each entry [i][j] counts how many actual rating-i predictions were made as j."),
]:
    code(doc3, line)
    body(doc3, expl, sa=4)

heading(doc3, "Section 6 — User Group Evaluation [NEW]")
code(doc3, "groups = {'Light': ..., 'Medium': ..., 'Heavy': ...}")
body(doc3,
    "Partitions training users into three activity groups. Light: 10–30 ratings "
    "(1,241 users). Medium: 31–100 (2,361 users). Heavy: >100 (2,438 users). "
    "No cold users (< 10) exist in MovieLens-1M — the dataset was filtered "
    "to include only active users.", sa=4)
code(doc3, "subset = test_df[test_df['user_id'].isin(uset)].head(GROUP_SAMPLES)")
body(doc3,
    "Takes up to 300 test samples from each group. Running both Temporal and "
    "TopK+Temporal on 900 total samples adds about 60 seconds of runtime.", sa=4)

heading(doc3, "Section 7 — Alpha Sensitivity Analysis [NEW]")
code(doc3, "alpha_values = [0.001, 0.01, 0.1, 0.5, 1.0]")
body(doc3,
    "Tests five orders of magnitude for Laplace smoothing. Very small α "
    "(0.001) trusts empirical frequencies; very large α (1.0) pulls all "
    "probabilities toward uniform. The alpha parameter is passed explicitly "
    "to ub_score() and ub_score_temporal() to override the global ALPHA.", sa=4)

heading(doc3, "Section 8 — Sustainability / Latency")
code(doc3, "avg_lat_htk = np.mean(latencies_htk) * 1000")
code(doc3, "p95_lat_htk = np.percentile(latencies_htk, 95) * 1000")
body(doc3,
    "Average and 95th-percentile latency in milliseconds. p95 is the "
    "tail latency — the worst-case time for 95% of predictions.", sa=4)

heading(doc3, "Section 9 — Output Explanations")
body(doc3,
    "All figures and tables are generated by Sections 5–8 of nbcf_final.py and written "
    "to the outputs/ directory. Each output is reproduced below with a technical "
    "interpretation of what it reveals about model behavior.")

heading(doc3, "Tables 1–5 — Algorithm Walkthrough", level=2)
body(doc3,
    "Tables 1–5 provide a step-by-step trace of the NBCF computation for a concrete "
    "user–item pair sampled from ML-1M. Table 1 renders a rating matrix excerpt that "
    "illustrates the 95.53% sparsity structure: missing entries (marked '•') dominate "
    "every row and column. Tables 2 and 3 display the Laplace-smoothed prior probabilities "
    "P(y|i) and user-based conditional likelihoods P(r_{u,j}=k | y, i) for the target pair. "
    "Likelihood values near 0.2 indicate that no co-occurrence observations exist between "
    "item j and rating class y, causing the estimate to revert toward the uniform baseline "
    "introduced by the α=0.01 smoothing term. Tables 4 and 5 present the analogous "
    "item-based quantities and the full posterior score comparison across all nine model "
    "variants; the asterisk (*) marks the MAP prediction for each variant.", sa=6)
insert_image(doc3, OUT + "table1_rating_matrix.png", width_cm=13)
caption(doc3, "Table 1. Rating matrix excerpt (ML-1M). Missing entries shown as '•'.")
insert_image(doc3, OUT + "table2_3_priors_likelihood.png", width_cm=13)
caption(doc3, "Tables 2–3. Laplace-smoothed prior probabilities and UB likelihoods for the running example.")
insert_image(doc3, OUT + "table4_5_ib_scores.png", width_cm=13)
caption(doc3, "Tables 4–5. IB likelihoods and full posterior score table for all nine model variants.")

heading(doc3, "Figure 1 — Confusion Matrices", level=2)
body(doc3,
    "The four panels characterize the classification behavior of the TopK+Temporal model. "
    "Panels (a) and (b) present the 5-class confusion matrix in raw counts and row-normalized "
    "form; panels (c) and (d) collapse the task to a binary Like (rating ≥ 4) / Dislike (rating < 4) "
    "decision. The normalized diagonal quantifies per-class recall. Systematic off-diagonal "
    "concentration in rows 1 and 2 reflects the modal-class bias inherent in NBCF: training "
    "evidence for rare classes is sparse, causing the model to predict the high-frequency "
    "classes (3, 4) more often than the ground truth warrants. The binary confusion matrix "
    "(panels c–d) shows substantially better discrimination, consistent with the higher "
    "binary F1 reported in Table 7.", sa=6)
insert_image(doc3, OUT + "figure1_confusion_matrices.png", width_cm=13)
caption(doc3, "Figure 1. Confusion matrices for TopK+Temporal: 5-class and binary (Like/Dislike).")

heading(doc3, "Figure 2 — MAE and RMSE Comparison", level=2)
body(doc3,
    "The grouped bar chart presents MAE (left panel) and RMSE (right panel) for all five model "
    "variants. RMSE bars consistently exceed their MAE counterparts because RMSE penalizes "
    "large errors quadratically, amplifying the contribution of outlier predictions. The "
    "pronounced gap between UB (MAE=0.861) and IB (MAE=2.335) demonstrates the noise "
    "introduced by uninformative co-raters on popular items — a problem that Top-K "
    "selection mitigates for UB but only partially resolves for the IB component. "
    "The near-identical bars for Temporal and TopK+Temporal confirm that the K=30 "
    "approximation preserves accuracy while reducing latency by 59%.", sa=6)
insert_image(doc3, OUT + "figure2_mae_rmse_comparison.png", width_cm=13)
caption(doc3, "Figure 2. MAE and RMSE comparison across all five NBCF variants.")

heading(doc3, "Figure 3 — Per-Class F1", level=2)
body(doc3,
    "The grouped bars represent per-class F1 for the Hybrid, Temporal, and TopK+Temporal "
    "variants across all five rating classes. F1 is highest for class 4 because it accounts "
    "for approximately 38% of ML-1M training ratings, providing the densest likelihood "
    "estimates for that class. Classes 1 and 2 yield near-zero F1 across all variants — "
    "a direct consequence of class imbalance combined with the Laplace prior pulling "
    "probability mass away from unseen rating classes rather than toward them. The marginal "
    "differences between Temporal and TopK+Temporal fall within the variance expected for "
    "a 1,000-sample evaluation.", sa=6)
insert_image(doc3, OUT + "figure3_per_class_f1.png", width_cm=13)
caption(doc3, "Figure 3. Per-class F1 scores for the three main NBCF variants.")

heading(doc3, "Figure 4 — Confidence Analysis", level=2)
body(doc3,
    "Panel (a) plots the distribution of max-probability confidence scores over the 1,000-sample "
    "test set. The right-skewed shape — peaked in the 0.30–0.55 range — reflects the intrinsic "
    "uncertainty of 5-class rating prediction; scores near 1.0 are rare and indicate that a "
    "single class strongly dominates the posterior. The dashed vertical line marks the operating "
    "threshold τ=0.45. Panel (b) plots mean MAE per confidence bucket, revealing a monotonic "
    "relationship: higher confidence corresponds to lower prediction error. This confirms that "
    "the max-probability measure functions as a meaningful calibrated uncertainty proxy and "
    "justifies its use as an abstention criterion.", sa=6)
insert_image(doc3, OUT + "figure4_confidence.png", width_cm=13)
caption(doc3, "Figure 4. Confidence score distribution (left) and MAE per confidence bucket (right).")

heading(doc3, "Figure 5 — Rating Distribution", level=2)
body(doc3,
    "Panel (a) compares the empirical rating distribution of the test set against the "
    "predicted distribution for all three main variants. Every model over-predicts class 4, "
    "mirroring the training distribution imbalance — the NBCF posterior is pulled toward "
    "the empirical mode rather than fitting the full distribution shape. Panel (b) plots "
    "per-class MAE, showing that extreme ground-truth classes (1 and 5) carry the highest "
    "individual errors. This is an expected consequence of the smoothed prior: when actual "
    "ratings are far from the modal class, the posterior's attraction to that mode inflates "
    "the absolute error.", sa=6)
insert_image(doc3, OUT + "figure5_rating_distribution.png", width_cm=13)
caption(doc3, "Figure 5. Actual vs predicted rating distribution (left) and per-class MAE (right).")

heading(doc3, "Figure 6 — Latency and Throughput", level=2)
body(doc3,
    "Panel (a) overlays per-prediction latency histograms for the Temporal and "
    "TopK+Temporal models across 1,000 predictions. The TopK+Temporal distribution "
    "is shifted left (lower median) and narrower (lower variance), reflecting the O(K) "
    "bound on the inner scoring loop. Panel (b) plots predictions per second against "
    "user history depth. Full Temporal throughput degrades as history grows because the "
    "loop iterates over all past items; TopK+Temporal throughput remains stable at K=30 "
    "operations per prediction regardless of total item count, decoupling latency from "
    "user activity level.", sa=6)
insert_image(doc3, OUT + "figure6_sustainability.png", width_cm=13)
caption(doc3, "Figure 6. Per-prediction latency distribution (left) and throughput vs history depth (right).")

heading(doc3, "Figure 7 — User Group Performance [NEW]", level=2)
body(doc3,
    "The bar charts decompose MAE (left) and RMSE (right) by user activity group for "
    "the Temporal and TopK+Temporal models. Light users (10–30 ratings) are unaffected "
    "by K=30 because their history never exceeds the cap. Medium users (31–100) benefit "
    "from Top-K neighborhood selection — a 6.2% MAE reduction — because the top-30 "
    "most-overlapping neighbors provide cleaner likelihood estimates than the full, "
    "unfiltered history. Heavy users (>100 ratings) incur a slight regression "
    "(MAE 0.740 → 0.777) because K=30 occasionally discards genuinely informative "
    "neighbors beyond rank 30, suggesting that adaptive K scaling with user activity "
    "level is a productive direction for future work.", sa=6)
insert_image(doc3, OUT + "figure7_user_groups.png", width_cm=13)
caption(doc3, "Figure 7. MAE and RMSE by user activity group (Light / Medium / Heavy).")

heading(doc3, "Figure 8 — Confidence Threshold Sweep [NEW]", level=2)
body(doc3,
    "The dual-axis chart plots coverage (fraction of predictions accepted; blue, left axis) "
    "and filtered MAE over accepted predictions (orange dashed, right axis) as the confidence "
    "threshold τ varies from 0.20 to 0.95. At the operating point τ=0.45, coverage is 29.3% "
    "and filtered MAE is 0.706 — an 11.5% improvement over the unfiltered MAE of 0.798. "
    "Both curves decrease monotonically as τ rises, confirming that the confidence measure "
    "is well-calibrated: higher-confidence predictions are systematically more accurate. "
    "The smooth trade-off curve allows the threshold to be tuned for specific deployment "
    "requirements balancing coverage against prediction quality.", sa=6)
insert_image(doc3, OUT + "figure8_confidence_sweep.png", width_cm=13)
caption(doc3, "Figure 8. Coverage–MAE trade-off as the confidence threshold τ is swept from 0.20 to 0.95.")

heading(doc3, "Figure 9 — Alpha Sensitivity [NEW]", level=2)
body(doc3,
    "The log-scaled x-axis spans five values of the Laplace smoothing parameter α. "
    "The UB (no temporal; blue) curve exhibits a U-shaped profile with a minimum near "
    "α=0.1 (MAE=0.845): both under-smoothing (α=0.001, overfitting sparse likelihoods) "
    "and over-smoothing (α=1.0, collapsing toward the uniform prior) degrade accuracy. "
    "The UB+Temporal (orange) curve is nearly flat across α ≥ 0.01 (MAE ≈ 0.855), "
    "demonstrating that exponential recency weighting absorbs the variance otherwise "
    "regularized by α, making the temporal model robust to the choice of smoothing "
    "hyperparameter. The dashed vertical line marks the default α=0.01.", sa=6)
insert_image(doc3, OUT + "figure9_alpha_sensitivity.png", width_cm=13)
caption(doc3, "Figure 9. MAE vs Laplace smoothing parameter α (log scale) for UB and UB+Temporal.")

heading(doc3, "Table 6 — User Group Metrics [NEW]", level=2)
body(doc3,
    "Numeric companion to Figure 7. Reports sample size n, MAE, and RMSE for each "
    "activity group under both the Temporal and TopK+Temporal models. The table enables "
    "precise quantification of the group-level performance differences described above. "
    "The monotonically decreasing MAE from Light to Heavy users quantifies the cold-start "
    "effect: additional historical ratings reduce prediction uncertainty by enriching the "
    "likelihood estimates with more co-occurrence evidence.", sa=6)
insert_image(doc3, OUT + "table6_user_groups.png", width_cm=13)
caption(doc3, "Table 6. MAE and RMSE broken down by user activity group.")

heading(doc3, "Table 7 — Full Metrics Summary (Updated)", level=2)
body(doc3,
    "Comprehensive evaluation table for all five NBCF variants. Columns follow the standard "
    "recommender systems reporting format: MAE, RMSE, macro precision, recall, and F1. "
    "Coverage is reported only for TopK+Temporal because the confidence-based abstention "
    "mechanism is applied exclusively to that variant. Precision, recall, and F1 are "
    "macro-averaged (equal weight per class), making them sensitive to performance on "
    "rare classes. The gap between macro F1 (0.27–0.37) and single-class accuracy on "
    "the modal class reflects the class-imbalance challenge described in the Figure 3 "
    "analysis.", sa=6)
insert_image(doc3, OUT + "table7_full_metrics.png", width_cm=13)
caption(doc3, "Table 7. Full metric comparison table for all five NBCF variants.")

doc3.save('Report_3_Final.docx')
print("✓ Report_3_Final.docx")


# ══════════════════════════════════════════════════════════════════════
# INFO_REPORT.TXT  —  Kapsamlı Değişiklik ve Rapor Özet Belgesi
# ══════════════════════════════════════════════════════════════════════

info = """
╔══════════════════════════════════════════════════════════════════════╗
║         NBCF FINAL PROJESİ — BİLGİ VE DEĞİŞİKLİK RAPORU           ║
║   Ömer MANAV (210209003) — Kaan Bahadır SAĞRA (210201011)          ║
║   Danışman: Gamze USLU, Ph.D.    Tarih: Haziran 2026               ║
╚══════════════════════════════════════════════════════════════════════╝

Bu belge üç amaca hizmet eder:
  1. Midterm ile final arasında neler değişti ve neden?
  2. Her değişiklik sistemi nasıl etkiledi?
  3. Report_1, Report_2 ve Report_3 ne içeriyor?

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
BÖLÜM 1 — PROJE NEDİR? (GENEL BAKIŞ)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Proje, Naïve Bayes Collaborative Filtering (NBCF) algoritmasını MovieLens-1M
veri seti üzerinde uygular. Sistem, bir kullanıcının henüz izlemediği bir
filme kaç puan vereceğini (1–5 arası) tahmin eder.

Temel fikir:
  - Kullanıcı u, film i'yi kaç puan verir?
  - Geçmişte u'nun izlediği filmler ile i'yi izleyen diğer kullanıcıların
    geçmişini karşılaştırarak Bayes teoremi ile olasılık hesapla.
  - En yüksek olasılıklı sınıfı tahmin olarak döndür.

Veri seti:
  - 6,040 kullanıcı, 3,701 film, 1,000,209 puan
  - %95.53 boşluk oranı (çoğu kullanıcı-film çifti için puan yok)
  - Zaman damgaları 2000–2003 yılları arası

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
BÖLÜM 2 — MİDTERM'DEN FİNALE DEĞİŞİKLİKLER
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

── DEĞİŞİKLİK 1: TOP-K KOMŞU SEÇİMİ ──────────────────────────────

  Midterm'de ne vardı:
    Her tahmin yapılırken kullanıcının geçmişteki TÜM filmleri (UB için)
    veya filmin TÜM birlikte puanlayıcıları (IB için) kullanılıyordu.
    Popüler bir film 2000+ kişi tarafından izlenmiş olabilir; bunların
    büyük çoğunluğu hedef kullanıcıyla ortak film izlememişse sıfır
    bilgi taşır ama yine de hesaplamaya katılır.

  Finalde ne değişti:
    Komşuları puanlamadan önce "co-rater overlap" ile sıralıyoruz:
    - UB için: |users(i) ∩ users(j)| — film i ve film j'yi aynı kişiler
      izlemişse j, i için güçlü kanıt demektir
    - IB için: |items(u) ∩ items(v)| — kullanıcı u ve v aynı filmleri
      izlemişse v, u için bilgilendirici demektir
    En yüksek örtüşmeye sahip K=30 komşu seçilir, gerisi atılır.

  Sisteme etkisi:
    + HIZLANMA: Tahmin süresi 34.8ms → 14.3ms (%59 düşüş)
      Sebebi: iç döngü sınırsız → K=30 ile sınırlı
    + DOĞRULUK (orta aktif kullanıcılar): MAE 0.803 → 0.753 (%6.2 iyileşme)
      Sebebi: alakasız komşular çıkarıldı, kalan 30 komşu daha bilgilendirici
    - DOĞRULUK (çok aktif kullanıcılar): MAE 0.740 → 0.777 (hafif kötüleşme)
      Sebebi: 100+ geçmişi olan kullanıcılarda K=30 bazen faydalı komşuları kesiyor
    ~ DOĞRULUK (az aktif kullanıcılar): değişmedi
      Sebebi: zaten 30'dan az geçmişleri var, K hiç devreye girmiyor

── DEĞİŞİKLİK 2: GÜVENİLİRLİK (CONFIDENCE) DÜZELTMESİ ──────────────

  Midterm'de ne vardı:
    Entropy tabanlı confidence hesabı: conf = 1 - H/log(5)
    Sonuç: %0 kapsama — tüm tahminlerin güven skoru eşik altında kalıyordu.

  Neden bozuluyordu:
    Hibrit birleştirme formülü: P_hybrid(y) = P_UB(y)^w_UB × P_IB(y)^w_IB
    Ağırlıklar: w_UB = 1/(1+25) ≈ 0.04 (25 geçmiş filmi olan kullanıcı için)
    Herhangi bir p değeri için p^0.04 ≈ 1.0 (sıfıra yakın üs → 1'e yakın sonuç)
    Tüm 5 sınıf için score ≈ 0.20 → tamamen uniform dağılım → maksimum entropy
    → confidence = 0

  Finalde ne değişti:
    Yeni formül: confidence = max_y P(y | u, i)
    Geometric birleştirme yapılmadan UB-Temporal skor üzerinde hesaplanır.
    Aralık: [0.2, 1.0] — 0.2 = tamamen belirsiz, 1.0 = kesin tahmin

  Sisteme etkisi:
    + Kapsama: %0 → %29.3 (τ=0.45 eşiğinde)
    + Yüksek güvenli tahminlerin MAE'si: 0.706 (genel 0.798'e göre %11.5 daha iyi)
    → Sistem artık "emin olmadığında" bunu doğru bir şekilde ifade edebiliyor

── DEĞİŞİKLİK 3: RMSE METRİĞİ ─────────────────────────────────────

  Midterm'de ne vardı: Sadece MAE
  Finalde ne değişti: MAE + RMSE birlikte hesaplanıyor

  RMSE = √( (1/n) × Σ(gerçek - tahmin)² )

  MAE ile farkı: Büyük hatalar kare alındığı için daha ağır cezalandırılır.
  MAE=0.79 iyi görünse de RMSE=1.12 gösteriyor ki bazı tahminler 2-3 puan
  yanılıyor. İkisini birlikte raporlamak daha eksiksiz bir resim verir.

  Sisteme etkisi: Algoritmayı değiştirmez; değerlendirmeyi daha kapsamlı yapar.

── DEĞİŞİKLİK 4: KULLANICI GRUBU DEĞERLENDİRMESİ ────────────────────

  Midterm'de ne vardı:
    Sadece 20-30 puan vermiş kullanıcılar test edildi.
    Bu, tüm kullanıcıların yalnızca %20'sini temsil ediyor.

  Finalde ne değişti:
    3 grup: Light (10-30), Medium (31-100), Heavy (>100)
    Her gruptan 300 örnek test edildi.

  Sisteme etkisi:
    Cold-start problemi görünür hale geldi:
    - Light kullanıcılar: MAE ≈ 0.84 (daha az geçmiş = daha zor tahmin)
    - Heavy kullanıcılar: MAE ≈ 0.74 (daha fazla geçmiş = daha kolay tahmin)
    Top-K'nın nerede fayda sağladığı, nerede zarar verdiği artık ölçülüyor.

── DEĞİŞİKLİK 5: ALPHA DUYARLILIK ANALİZİ ──────────────────────────

  Midterm'de ne vardı: α=0.01 sabit, hiç sorgulanmadı.
  Finalde ne değişti: α ∈ {0.001, 0.01, 0.1, 0.5, 1.0} için MAE hesaplandı.

  Bulgular:
    - UB (temporal yok): α=0.1 civarında en düşük MAE (0.845)
    - UB+Temporal: α değişince MAE neredeyse hiç değişmiyor (0.855 sabit)
    → Temporal weighting, smoothing parametresine karşı modeli sağlamlaştırıyor

── GÖRSEL ÇIKTILAR ─────────────────────────────────────────────────

  Midterm:  6 figür, 5 tablo
  Final:    9 figür, 7 tablo

  Yeni eklenenler:
    Figure 7 — Kullanıcı grubu MAE/RMSE karşılaştırması
    Figure 8 — Güven eşiği sweep'i (coverage vs filtered MAE)
    Figure 9 — Alpha duyarlılık grafiği
    Table 6  — Kullanıcı grubu sayısal özet
    Table 7  — Güncellendi: RMSE kolonu + TopK+Temporal satırı eklendi
    Figure 2 — Güncellendi: artık RMSE de gösteriyor
    Figure 3 — Güncellendi: 3. model (TopK) eklendi

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
BÖLÜM 3 — METRIK KARŞILAŞTIRMASI (ÖZET TABLO)
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

  Model                  MAE      RMSE     F1      Hız(ms)  Coverage
  ─────────────────────────────────────────────────────────────────────
  UB (baseline)          0.8610   1.2518   —       —        —
  IB (baseline)          2.3350   2.6384   —       —        —
  Hybrid (baseline)      1.1020   1.4953   0.271   28.3ms   —
  Hybrid + Temporal      0.7920   1.1234   0.279   34.8ms   0% (BOZUK)
  TopK + Temporal [YENİ] 0.7980   1.1261   0.266   14.3ms   29.3%

  En önemli gelişme: Hız (%59 düşüş) + Confidence onarımı
  En kritik bulgu: Temporal modelin IB'yi düzeltemediği, UB'nin
  dominant kaldığı ve Top-K'nın sadece orta aktif kullanıcılarda
  doğruluk kazanımı sağladığı

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
BÖLÜM 4 — REPORT İÇERİKLERİ
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

── REPORT_1_Final.docx ─────────────────────────────────────────────
  Tam akademik makale formatında. Hocanın istediği 8 bölüm aynen mevcut:

  Abstract       — 150 kelimelik özet: veri seti, yöntem, ana sonuçlar
  Introduction   — Motivasyon, veri seti detayları, formal problem tanımı,
                   zorluklar (sparsity, temporal drift, scalability, confidence)
  Related Work   — 7 kaynak: Breese (1998), Si&Jin (2003), Koren (2009),
                   Sarwar (2001), Herlocker (2004), Harper (2015), Olmo (2008)
  Proposed       — NBCF matematigi, log-space, temporal decay, Top-K seçimi,
    Approach       hibrit birleştirme, confidence skorlama
  Experimental   — Train/test split, evaluation subsets, hyperparameters,
    Setup          metrik tanımları
  Performance    — Genel sonuçlar + tablo, kullanıcı grubu analizi,
    Analysis       confidence analizi, alpha duyarlılığı, latency
  Conclusion     — Kazanımlar, sınırlılıklar, gelecek çalışmalar
  References     — 7 akademik kaynak düzgün format ile

  Görseller: Tüm 9 figür ve 7 tablo ilgili bölümlere yerleştirildi.
  Sayfa sayısı: ~20-25 sayfa (görsellerle birlikte)

── REPORT_2_Final.docx ─────────────────────────────────────────────
  Hocanın istediği: "midterm ile final arasındaki farkların 1 sayfalık özeti"

  İçerik:
  1. Overview — Değişikliklerin genel çerçevesi (2 paragraf)
  2. Algorithm Changes
     2.1 Top-K Neighborhood Selection — problem, çözüm, etki
     2.2 Confidence Score Fixed — neden bozuktu, ne değişti, etkisi
  3. Evaluation Changes
     3.1 RMSE metriği
     3.2 Kullanıcı grubu değerlendirmesi
     3.3 Alpha duyarlılık analizi
     3.4 Confidence threshold sweep
  4. Metric Comparison Table — midterm vs final sayısal karşılaştırma
  5. Output Changes Table — figür/tablo sayısı, model sayısı, metrikler

  Not: 1 sayfa olması isteniyor ama içerik biraz daha fazla çıktı.
       Yazı boyutunu küçülterek veya bölümleri kısaltarak 1 sayfaya
       sığdırabilirsiniz.

── REPORT_3_Final.docx ─────────────────────────────────────────────
  Hocanın istediği: "kodu satır satır ve bölüm bölüm açıkla +
                     her çıktı bölümünü açıkla"

  İçerik:
  Overview       — Scriptin 9 bölümü hakkında genel açıklama
  Section 1      — Import kütüphaneleri satır satır + global sabitler
  Section 2      — Veri yükleme, train/test split, sözlük yapısı,
                   frozenset precomputation [YENİ açıklandı]
  Section 3      — Tüm orijinal fonksiyonlar: prior, likelihood,
                   normalize, ub_score, ib_score, temporal fonksiyonlar
  Section 4      — [YENİ] topk_items_for_ub, topk_users_for_ib,
                   confidence_maxprob, hybrid_topk_temporal
  Section 5      — Ana değerlendirme döngüsü: RMSE [YENİ], confidence_max [YENİ]
  Section 6      — [YENİ] Kullanıcı grubu değerlendirmesi
  Section 7      — [YENİ] Alpha duyarlılık analizi
  Section 8      — Latency ölçümü
  Section 9      — Her çıktının açıklaması:
                   Tables 1-5, Figures 1-9, Table 6, Table 7

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
BÖLÜM 5 — DOSYA YAPISI
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

  report/final/
  ├── nbcf_final.py              — Ana kod dosyası
  ├── generate_reports.py        — Raporları üreten script (teslim edilmeyecek)
  ├── Report_1_Final.docx        — Akademik makale (görseller dahil)
  ├── Report_2_Final.docx        — Midterm vs final farkları
  ├── Report_3_Final.docx        — Satır satır kod açıklaması
  ├── INFO_REPORT.txt            — Bu belge (teslim edilmeyecek, referans)
  ├── README.md                  — Kurulum ve çalıştırma adımları (HAZIRLANACAK)
  ├── ml-1m/                     — Dataset
  │   ├── ratings.dat
  │   ├── movies.dat
  │   └── users.dat
  └── outputs/                   — Üretilen görsel çıktılar
      ├── figure1_confusion_matrices.png
      ├── figure2_mae_rmse_comparison.png
      ├── figure3_per_class_f1.png
      ├── figure4_confidence.png
      ├── figure5_rating_distribution.png
      ├── figure6_sustainability.png
      ├── figure7_user_groups.png      [YENİ]
      ├── figure8_confidence_sweep.png [YENİ]
      ├── figure9_alpha_sensitivity.png [YENİ]
      ├── table1_rating_matrix.png
      ├── table2_3_priors_likelihood.png
      ├── table4_5_ib_scores.png
      ├── table6_user_groups.png       [YENİ]
      └── table7_full_metrics.png

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
BÖLÜM 6 — TEKNİK NOTLAR
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

Neden IB hâlâ bu kadar kötü (MAE=2.33)?
  IB, item i'nin TÜM birlikte puanlayıcılarını kullanır. Popüler filmler
  için bu 2000+ kişidir. Çoğu kişi hedef kullanıcıyla hiç ortak film
  izlememiştir → likelihood = alpha/(NUM_R×alpha) = 0.2 (uniform).
  500 kişilik uniform likelihood log toplamı = 500 × log(0.2) = -800,
  bu tüm sınıflar için eşit, dolayısıyla prior domine eder.
  Prior çoğunlukla 4 veya 5 yönünde → model 4-5 tahmin eder →
  1-3 puanlı filmler için büyük hata. Top-K IB bunu kısmen düzeltiyor
  ama tam çözüm için Jaccard similarity veya cosine similarity gerekir.

Neden eval seti küçük (1000)?
  Her tahmin 14-35ms sürer. 200,000 test örneğinin tamamı ~80 dakika
  isterdi. 1000 örnek ~2 dakika — pratikte makul bir denge.

Temporal Lambda neden 1e-7?
  ML-1M timestamps saniye cinsinden. 1 yıl ≈ 3.15×10^7 saniye.
  λ=1e-7 ile: 1 yıllık rating ağırlığı = exp(-1e-7 × 3.15e7) ≈ 0.04
  3 yıllık rating: exp(-1e-7 × 9.5e7) ≈ 0.09
  Yani 1-3 yıllık ratingler ciddi şekilde ağırlığını kaybediyor.
"""

with open('INFO_REPORT.txt', 'w', encoding='utf-8') as f:
    f.write(info.strip())
print("✓ INFO_REPORT.txt")

print("\nTüm dosyalar hazır:", os.getcwd())
